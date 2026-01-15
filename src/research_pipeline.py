import os
import re
import logging
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# --- Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Configuration ---
PAPER_DIR = "paper"
INPUT_FILE = os.path.join(PAPER_DIR, "paper.md")
OUTPUT_DOCX = os.path.join(PAPER_DIR, "research_paper.docx")
OUTPUT_PDF = os.path.join(PAPER_DIR, "research_paper.pdf")


def set_2_columns(section):
    """Sets a section to have 2 columns using OXML."""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '450')  # 450 twips = ~0.3125 inches (0.8cm)
    if not sectPr.xpath('./w:cols'):
        sectPr.append(cols)


def configure_styles(doc):
    """Sets strict styling for Normal, Heading 1, etc."""
    styles = doc.styles

    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    p_format = style_normal.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_format.first_line_indent = Inches(0.17)
    p_format.space_after = Pt(0)


def parse_inline_formatting(paragraph, text, font_size=Pt(10), base_bold=False, base_italic=False):
    """
    Parses inline markdown (**bold**, *italic*) and adds runs to the paragraph.
    Uses a non-greedy regex to correctly identify formatting tokens.
    """
    # Pattern: Match **bold**, *italic*, or plain text segments
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')
    
    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            plain_text = text[last_end:match.start()]
            run = paragraph.add_run(plain_text)
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.font.bold = base_bold
            run.font.italic = base_italic

        matched_text = match.group(0)
        if matched_text.startswith('**') and matched_text.endswith('**'):
            # Bold text
            run = paragraph.add_run(matched_text[2:-2])
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.font.bold = True
            run.font.italic = base_italic
        elif matched_text.startswith('*') and matched_text.endswith('*'):
            # Italic text
            run = paragraph.add_run(matched_text[1:-1])
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.font.bold = base_bold
            run.font.italic = True
        
        last_end = match.end()

    # Add remaining plain text after last match
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = 'Times New Roman'
        run.font.size = font_size
        run.font.bold = base_bold
        run.font.italic = base_italic


def process_markdown(doc, md_content):
    """Converts markdown content to a Word document with IEEE formatting."""
    lines = md_content.split('\n')

    current_section_type = "header"
    in_references = False
    in_table = False
    table_buffer = []

    # Setup Page 1 Margins (Standard US Letter)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.125)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    for line in lines:
        line = line.strip()

        # --- TABLE PARSING ---
        if line.startswith('|'):
            in_table = True
            row_content = [c.strip() for c in line.strip('|').split('|')]
            if '---' in row_content[0]:
                continue
            table_buffer.append(row_content)
            continue
        elif in_table:
            # Flush table
            if table_buffer:
                table = doc.add_table(rows=len(table_buffer), cols=len(table_buffer[0]))
                table.style = 'Table Grid'
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_idx, row_data in enumerate(table_buffer):
                    row = table.rows[r_idx]
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.first_line_indent = Pt(0)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            run = p.add_run(cell_text)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)  # Match body text size per IEEE spec
                            if r_idx == 0:
                                run.font.bold = True
                doc.add_paragraph()
            table_buffer = []
            in_table = False

        if not line:
            continue

        # --- DETECT REFERENCES SECTION ---
        if line.upper().startswith('## REFERENCES') or line == '## REFERENCES':
            in_references = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("REFERENCES")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.small_caps = True
            continue

        # --- SWITCH TO 2-COLUMN BODY ---
        is_body_trigger = (
            line.startswith('**Abstract**') or
            line.startswith('**Index Terms**') or
            line.startswith('I. ')
        )

        if current_section_type == "header" and is_body_trigger:
            new_section = doc.add_section(WD_SECTION.CONTINUOUS)
            set_2_columns(new_section)
            current_section_type = "body"

        # --- PARSING LOGIC ---

        if line.startswith('# '):
            # MAIN TITLE: 24pt Centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(24)
            run.font.bold = False

        elif line.startswith('% '):
            # AUTHOR/AFFILIATION
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            is_affil = '@' in line or 'Department' in line or 'Institute' in line
            run = p.add_run(line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.italic = is_affil

        elif line.startswith('**Abstract**'):
            # Abstract - Bold Italic
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            clean_line = line.replace('**Abstract**', 'Abstract').replace('**', '').replace('*', '')
            run = p.add_run(clean_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.italic = True

        elif line.startswith('**Index Terms**'):
            # Index Terms - Bold Italic
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            clean_line = line.replace('**Index Terms**', 'Index Terms').replace('**', '').replace('*', '')
            run = p.add_run(clean_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.italic = True

        elif line.startswith('## '):
            # LEVEL 1 HEADING: Small Caps, Centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            text = line[3:].upper()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.small_caps = True

        elif line.startswith('### '):
            # LEVEL 2 HEADING: Italic, Left Aligned
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[4:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True

        elif line.upper().startswith('TABLE '):
            # TABLE CAPTION: Small Caps, Left Justified
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line.upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.small_caps = True

        elif re.match(r'^\d+\)', line):
            # NUMBERED LIST ITEM: e.g., "1) Item text"
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            font_size = Pt(9) if in_references else Pt(10)
            parse_inline_formatting(p, line, font_size=font_size)

        elif re.match(r'^\[\d+\]', line):
            # REFERENCE ENTRY: e.g., "[1] Author..."
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            parse_inline_formatting(p, line, font_size=Pt(9))

        elif line.startswith('* ') or line.startswith('- '):
            # BULLET LIST ITEM
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Inches(0.5)
            bullet_text = line[2:]
            run = p.add_run('• ')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            parse_inline_formatting(p, bullet_text, font_size=Pt(10))

        elif re.match(r'^!\[.*\]$', line):
            # IMAGE: Format ![filename.png]
            img_match = re.match(r'^!\[(.+)\]$', line)
            if img_match:
                img_filename = img_match.group(1)
                img_path = os.path.join(PAPER_DIR, img_filename)
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run()
                    # Width for single column: ~3.25 inches
                    run.add_picture(img_path, width=Inches(3.0))
                else:
                    logger.warning(f"Image not found: {img_path}")

        elif line.startswith('Figure ') and '.' in line:
            # FIGURE CAPTION: Small Caps, Centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.small_caps = True

        elif line.startswith('$$') and line.endswith('$$'):
            # DISPLAY EQUATION: Centered, Italic
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            eq_text = line[2:-2].strip()
            run = p.add_run(eq_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True

        else:
            # BODY TEXT with inline formatting
            p = doc.add_paragraph()
            font_size = Pt(9) if in_references else Pt(10)
            parse_inline_formatting(p, line, font_size=font_size)


def main():
    if not os.path.exists(INPUT_FILE):
        logger.error(f"Input file not found: {INPUT_FILE}")
        return

    try:
        with open(INPUT_FILE, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except IOError as e:
        logger.error(f"Failed to read input file: {e}")
        return

    doc = Document()
    configure_styles(doc)
    process_markdown(doc, md_content)

    try:
        doc.save(OUTPUT_DOCX)
        logger.info(f"Generated DOCX: {OUTPUT_DOCX}")
    except IOError as e:
        logger.error(f"Failed to save DOCX: {e}")
        return

    # Convert DOCX to PDF using LibreOffice
    try:
        abs_docx = os.path.abspath(OUTPUT_DOCX)
        abs_outdir = os.path.abspath(PAPER_DIR)
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', abs_outdir, abs_docx],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            logger.info(f"Generated PDF: {OUTPUT_PDF}")
        else:
            logger.error(f"LibreOffice conversion failed: {result.stderr}")
    except subprocess.TimeoutExpired:
        logger.error("PDF conversion timed out")
    except FileNotFoundError:
        logger.error("LibreOffice (soffice) not found. Install with: brew install --cask libreoffice")
    except Exception as e:
        logger.error(f"Failed to convert to PDF: {e}")


if __name__ == "__main__":
    main()